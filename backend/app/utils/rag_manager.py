# app/utils/rag_manager.py

import os
import re
import sqlite3
try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
from typing import List, Dict, Any, Optional
import numpy as np

DB_PATH = os.getenv("AUTOSOIL_SQLITE_PATH", "data.db")
if os.name != 'nt' and (':' in DB_PATH or '\\' in DB_PATH):
    DB_PATH = "data.db"

def _split_paths(value: str) -> List[str]:
    return [path.strip() for path in re.split(r"[;|]", value) if path.strip()]


GEOLOGS_DIR = os.getenv("AUTOSOIL_GEOLOGS_DIR", os.path.join(os.getcwd(), "Project Geologs"))
REPORTS_DIRS = _split_paths(os.getenv("AUTOSOIL_REPORTS_DIRS", "")) or [
    os.path.join(GEOLOGS_DIR, "03 - Reports"),
    os.path.join(GEOLOGS_DIR, "Reports 2"),
    os.path.join(GEOLOGS_DIR, "Reports 3")
]

# We will check if easyocr is available
try:
    import easyocr
    # Initialize reader once for English
    # easyocr might download model files on first init
    OCR_READER = easyocr.Reader(['en'], gpu=False) 
    print("[RAG Manager] EasyOCR loaded successfully.")
except Exception as e:
    OCR_READER = None
    print(f"[RAG Manager] EasyOCR failed to load: {e}. OCR will fall back to text-only extraction.")


def init_rag_db():
    """Initialize SQLite tables for RAG index."""
    os.makedirs(os.path.dirname(DB_PATH) if os.path.dirname(DB_PATH) else ".", exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # RAG index for past reports
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS reports_rag (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_path TEXT UNIQUE,
            file_name TEXT,
            group_name TEXT,
            sub_folder TEXT,
            content TEXT,
            category TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    # Cache for OCR results to avoid re-running slow OCR
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS ocr_cache (
            file_path TEXT PRIMARY KEY,
            extracted_text TEXT,
            metadata TEXT,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    conn.commit()
    conn.close()


def extract_docx_text(file_path: str) -> str:
    """Extract text from a .docx file paragraphs and tables."""
    if docx is None:
        print("[RAG Manager] python-docx is not installed. docx extraction skipped.")
        return ""
    if not os.path.exists(file_path):
        return ""
    try:
        doc = docx.Document(file_path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    except Exception as e:
        print(f"[RAG Manager] Error reading docx {file_path}: {e}")
        return ""


def extract_pdf_text(file_path: str, allow_ocr: bool = True) -> str:
    """Extract text from a .pdf file, with OCR fallback if needed."""
    if not os.path.exists(file_path):
        return ""
        
    # Check cache first
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT extracted_text FROM ocr_cache WHERE file_path = ?", (file_path,))
    cached = cursor.fetchone()
    conn.close()
    if cached:
        return cached[0]
        
    file_size = os.path.getsize(file_path)
    is_large = file_size > 5 * 1024 * 1024
    if is_large:
        print(f"[RAG Manager] PDF file is large ({file_size/(1024*1024):.2f} MB). Disabling OCR and limiting pages.")
        allow_ocr = False

    text = ""
    if pdfplumber is None:
        print("[RAG Manager] pdfplumber is not installed. Fast text extraction skipped.")
    else:
        try:
            # First try fast text extraction
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                pages_to_read = pdf.pages[:5] if is_large else pdf.pages
                for page in pages_to_read:
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
                text = "\n".join(pages_text).strip()
        except Exception as e:
            print(f"[RAG Manager] PDFPlumber error on {file_path}: {e}")

    # Fallback to pypdf if pdfplumber fails
    if not text:
        if PdfReader is None:
            print("[RAG Manager] pypdf is not installed. Fallback PDF text extraction skipped.")
        else:
            try:
                reader = PdfReader(file_path)
                pages_text = []
                pages_to_read = reader.pages[:5] if is_large else reader.pages
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
                text = "\n".join(pages_text).strip()
            except Exception as e:
                print(f"[RAG Manager] PyPDF error on {file_path}: {e}")

    # If it is a scanned document (no text found) and OCR reader is available, we try OCR
    if not text and allow_ocr and OCR_READER:
        print(f"[RAG Manager] PDF contains no text. Attempting OCR on {file_path}...")
        try:
            # We use pdfplumber's page.to_image() which doesn't require poppler binaries if we can draw or save it
            if pdfplumber is None:
                print("[RAG Manager] pdfplumber is not installed. PDF OCR extraction skipped.")
            else:
                with pdfplumber.open(file_path) as pdf:
                    ocr_pages = []
                    for i, page in enumerate(pdf.pages[:5]): # Limit to first 5 pages for speed
                        try:
                            img = page.to_image(resolution=150)
                            # Save temp image
                            temp_img_path = f"temp_page_{i}.png"
                            img.save(temp_img_path)
                            # Run OCR
                            result = OCR_READER.readtext(temp_img_path, detail=0)
                            if result:
                                ocr_pages.append("\n".join(result))
                            # Cleanup temp file
                            if os.path.exists(temp_img_path):
                                os.remove(temp_img_path)
                        except Exception as img_err:
                            print(f"[RAG Manager] PDF page image conversion/OCR failed: {img_err}")
                    text = "\n".join(ocr_pages).strip()
        except Exception as ocr_err:
            print(f"[RAG Manager] Scanned PDF OCR failed for {file_path}: {ocr_err}")

    # Cache it
    if text:
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            cursor.execute("INSERT OR REPLACE INTO ocr_cache (file_path, extracted_text) VALUES (?, ?)", (file_path, text))
            conn.commit()
            conn.close()
        except Exception as cache_err:
            print(f"[RAG Manager] Failed to cache OCR text: {cache_err}")

    return text


def extract_xlsx_text(file_path: str) -> str:
    """Extract text from an Excel file."""
    if not os.path.exists(file_path):
        return ""
    if openpyxl is None:
        print("[RAG Manager] openpyxl is not installed. xlsx extraction skipped.")
        return ""
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text_parts.append(f"--- Sheet: {sheet} ---")
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(cell).strip() for cell in row if cell is not None]
                if row_vals:
                    text_parts.append(" | ".join(row_vals))
        return "\n".join(text_parts)
    except Exception as e:
        print(f"[RAG Manager] Error reading xlsx {file_path}: {e}")
        return ""


def extract_image_ocr(file_path: str) -> str:
    """Perform OCR on an image file using EasyOCR."""
    if not os.path.exists(file_path):
        return ""
        
    # Check cache first
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT extracted_text FROM ocr_cache WHERE file_path = ?", (file_path,))
    cached = cursor.fetchone()
    conn.close()
    if cached:
        return cached[0]
        
    text = ""
    if OCR_READER:
        print(f"[RAG Manager] OCR-ing image: {file_path}...")
        try:
            result = OCR_READER.readtext(file_path, detail=0)
            text = "\n".join(result).strip()
        except Exception as e:
            print(f"[RAG Manager] OCR failed on image {file_path}: {e}")
            text = f"[OCR Failed: {e}]"
    else:
        text = "[OCR Unavailable: EasyOCR not installed]"
        
    # Cache it
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO ocr_cache (file_path, extracted_text) VALUES (?, ?)", (file_path, text))
    conn.commit()
    conn.close()
    
    return text


def build_rag_index(limit: Optional[int] = None):
    """Scan all historical reports and build the index database."""
    init_rag_db()
    print("[RAG Manager] Starting RAG index build...", flush=True)
    
    # 1. Fetch already indexed file paths
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT file_path FROM reports_rag")
    indexed_files = {row[0] for row in cursor.fetchall()}
    conn.close()
    
    indexed_count = 0
    
    # 2. Scan directories and extract content
    for base_dir in REPORTS_DIRS:
        if not os.path.exists(base_dir):
            print(f"[RAG Manager] Reports folder not found: {base_dir}", flush=True)
            continue
            
        group_name = os.path.basename(base_dir)
        print(f"[RAG Manager] Scanning group {group_name}...", flush=True)
        
        for root, dirs, files in os.walk(base_dir):
            for file in files:
                if file.startswith("~$") or not (file.endswith(".docx") or file.endswith(".pdf")):
                    continue
                    
                full_path = os.path.normpath(os.path.join(root, file))
                
                if full_path in indexed_files:
                    continue
                    
                # Skip massive files (drawings, maps, dwg/pdf sets) during bulk indexing
                if os.path.exists(full_path) and os.path.getsize(full_path) > 5 * 1024 * 1024:
                    print(f"[RAG Manager] Skipping large file during bulk indexing: {file} ({os.path.getsize(full_path)/(1024*1024):.2f} MB)", flush=True)
                    continue

                print(f"[RAG Manager] Indexing file {file}...", flush=True)
                sub_folder = os.path.basename(root) if root != base_dir else "General"
                content = ""
                
                if file.endswith(".docx"):
                    content = extract_docx_text(full_path)
                elif file.endswith(".pdf"):
                    content = extract_pdf_text(full_path, allow_ocr=False) # Disallow OCR in bulk index
                    
                if content.strip():
                    try:
                        conn = sqlite3.connect(DB_PATH)
                        cursor = conn.cursor()
                        cursor.execute(
                            "INSERT OR REPLACE INTO reports_rag (file_path, file_name, group_name, sub_folder, content, category) VALUES (?, ?, ?, ?, ?, ?)",
                            (full_path, file, group_name, sub_folder, content, sub_folder)
                        )
                        conn.commit()
                        conn.close()
                        indexed_count += 1
                        print(f"[RAG Manager] Indexed {file} successfully ({indexed_count}).", flush=True)
                    except Exception as db_err:
                        print(f"[RAG Manager] DB insert failed for {file}: {db_err}", flush=True)
                else:
                    print(f"[RAG Manager] Skipped empty/unreadable file: {file}", flush=True)
                
                if limit is not None and indexed_count >= limit:
                    print(f"[RAG Manager] Reached limit of {limit} indexed files. Stopping.", flush=True)
                    return

    print(f"[RAG Manager] RAG index updated. Added {indexed_count} new files.", flush=True)


def search_rag(query: str, limit: int = 5) -> List[Dict[str, Any]]:
    """Search for relevant snippets or documents based on matching terms."""
    init_rag_db()
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # We clean query and split into keywords
    keywords = re.findall(r'\w+', query.lower())
    if not keywords:
        cursor.execute("SELECT file_path, file_name, group_name, sub_folder, content FROM reports_rag LIMIT ?", (limit,))
        rows = cursor.fetchall()
        conn.close()
        return [dict(row) for row in rows]
        
    # We perform a simple LIKE search weighting hits for each keyword
    query_parts = []
    params = []
    for kw in keywords:
        query_parts.append("content LIKE ?")
        params.append(f"%{kw}%")
        
    sql = f"""
        SELECT file_path, file_name, group_name, sub_folder, content
        FROM reports_rag
        WHERE {" OR ".join(query_parts)}
        LIMIT 20
    """
    
    cursor.execute(sql, params)
    rows = cursor.fetchall()
    
    # Score relevance based on occurrences of keywords
    scored_results = []
    for row in rows:
        content_lower = row["content"].lower()
        score = sum(content_lower.count(kw) for kw in keywords)
        
        # Grab a snippet around the first keyword match
        snippet = ""
        first_match_idx = -1
        for kw in keywords:
            idx = content_lower.find(kw)
            if idx != -1:
                first_match_idx = idx
                break
                
        if first_match_idx != -1:
            start = max(0, first_match_idx - 150)
            end = min(len(row["content"]), first_match_idx + 250)
            snippet = row["content"][start:end].replace("\n", " ")
            if start > 0:
                snippet = "..." + snippet
            if end < len(row["content"]):
                snippet = snippet + "..."
        else:
            snippet = row["content"][:300] + "..."
            
        scored_results.append({
            "file_path": row["file_path"],
            "file_name": row["file_name"],
            "group_name": row["group_name"],
            "sub_folder": row["sub_folder"],
            "score": score,
            "snippet": snippet
        })
        
    conn.close()
    # Sort by score descending
    scored_results.sort(key=lambda x: x["score"], reverse=True)
    return scored_results[:limit]


def analyze_project_folder(folder_path: str) -> Dict[str, Any]:
    """Scan and analyze a specific site project directory to collect photos, logs, DCPs and extract text/tables."""
    if not os.path.exists(folder_path):
        return {"error": f"Folder {folder_path} not found"}
        
    analysis = {
        "folder_path": folder_path,
        "folder_name": os.path.basename(folder_path),
        "photos": [],
        "dcp_forms": [],
        "soil_logs": [],
        "other_documents": [],
        "summary": "",
        "bearing_capacity": "",
        "client": "",
        "site_address": "",
        "job_no": ""
    }
    
    print(f"[RAG Manager] Analyzing project folder: {folder_path}")
    
    # Walk the directory
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.startswith("~$"):
                continue
                
            full_path = os.path.normpath(os.path.join(root, file))
            file_lower = file.lower()
            
            file_info = {
                "file_name": file,
                "file_path": full_path,
                "size_bytes": os.path.getsize(full_path),
                "extracted_text": ""
            }
            
            # Categorize and extract text
            if any(ext in file_lower for ext in [".jpg", ".jpeg", ".png"]):
                file_info["extracted_text"] = extract_image_ocr(full_path)
                analysis["photos"].append(file_info)
            elif file_lower.endswith(".pdf"):
                extracted = extract_pdf_text(full_path)
                file_info["extracted_text"] = extracted
                # Detect category based on content/name
                if "dcp" in file_lower or "penetrometer" in extracted.lower():
                    analysis["dcp_forms"].append(file_info)
                elif "borehole" in file_lower or "soil log" in file_lower or "borehole log" in extracted.lower() or "soil profile" in extracted.lower():
                    analysis["soil_logs"].append(file_info)
                else:
                    analysis["other_documents"].append(file_info)
            elif file_lower.endswith(".docx"):
                extracted = extract_docx_text(full_path)
                file_info["extracted_text"] = extracted
                if "dcp" in file_lower or "penetrometer" in extracted.lower():
                    analysis["dcp_forms"].append(file_info)
                elif "borehole" in file_lower or "soil log" in file_lower or "borehole log" in extracted.lower():
                    analysis["soil_logs"].append(file_info)
                else:
                    analysis["other_documents"].append(file_info)
            elif file_lower.endswith(".xlsx") or file_lower.endswith(".xls"):
                extracted = extract_xlsx_text(full_path)
                file_info["extracted_text"] = extracted
                if "dcp" in file_lower or "penetrom" in file_lower:
                    analysis["dcp_forms"].append(file_info)
                elif "bore" in file_lower or "soil" in file_lower:
                    analysis["soil_logs"].append(file_info)
                else:
                    analysis["other_documents"].append(file_info)
                    
    # Generate simple summary of discoveries
    summary_parts = [
        f"Analyzed folder: {analysis['folder_name']}.",
        f"Found {len(analysis['photos'])} photos, {len(analysis['dcp_forms'])} DCP files, {len(analysis['soil_logs'])} Soil Logs, and {len(analysis['other_documents'])} other documents."
    ]
    
    # Try to extract key details (e.g. bearing capacity, clients, address, etc.)
    all_combined_text = ""
    for category in ["dcp_forms", "soil_logs", "other_documents", "photos"]:
        for f in analysis[category]:
            all_combined_text += f"\n{f['extracted_text']}"
            
    # Search for bearing capacity
    bc_match = re.search(r'(?:bearing capacity|allowable bearing|bearing)\s*:\s*([^\n\r]+)', all_combined_text, re.IGNORECASE)
    if bc_match:
        analysis["bearing_capacity"] = bc_match.group(1).strip()
        summary_parts.append(f"Detected Bearing Capacity: {analysis['bearing_capacity']}.")
        
    client_match = re.search(r'(?:Client|Commissioned By|Prepared For)\s*:\s*([^\n\r]+)', all_combined_text, re.IGNORECASE)
    if client_match:
        analysis["client"] = client_match.group(1).strip()
        summary_parts.append(f"Detected Client: {analysis['client']}.")
        
    address_match = re.search(r'(?:Site Address|Site|Location|Project Site|Address)\s*:\s*([^\n\r]+)', all_combined_text, re.IGNORECASE)
    if address_match:
        analysis["site_address"] = address_match.group(1).strip()
        summary_parts.append(f"Detected Site Address: {analysis['site_address']}.")
        
    job_match = re.search(r'(?:Job No|Job Number|Project No|Ref)\s*:\s*([A-Z0-9\-\s\/]+)', all_combined_text, re.IGNORECASE)
    if job_match:
        analysis["job_no"] = job_match.group(1).strip()
        summary_parts.append(f"Detected Job Number: {analysis['job_no']}.")

    analysis["summary"] = " ".join(summary_parts)
    return analysis


# ----------------------------------------------------
# Late Interaction (ColBERT-style) Search Mechanism
# ----------------------------------------------------

_LATE_INTERACTION_MODEL = None


def init_late_interaction_db():
    """Initialize SQLite tables for ColBERT late interaction."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS reports_late_interaction (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_path TEXT,
            chunk_index INTEGER,
            chunk_text TEXT,
            token_embeddings_blob BLOB,
            num_tokens INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    # Add index on file_path and chunk_index to speed up lookups or checks
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_late_interaction_file ON reports_late_interaction (file_path)")
    conn.commit()
    conn.close()


def chunk_text(text: str, chunk_size: int = 250, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks of words."""
    if not text:
        return []
    words = text.split()
    if not words:
        return []
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i : i + chunk_size]
        chunks.append(" ".join(chunk_words))
        if len(words) - i <= chunk_size:
            break
        i += chunk_size - overlap
    return chunks


def get_late_interaction_model():
    """Lazy load SentenceTransformer for token-level embeddings."""
    global _LATE_INTERACTION_MODEL
    if _LATE_INTERACTION_MODEL is None:
        from sentence_transformers import SentenceTransformer
        print("[RAG Manager] Loading SentenceTransformer 'all-MiniLM-L6-v2' for Late Interaction...", flush=True)
        _LATE_INTERACTION_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
    return _LATE_INTERACTION_MODEL


def build_index_from_reports(limit: Optional[int] = None):
    """Build the regular RAG index first, then construct ColBERT late interaction token embeddings index."""
    # Ensure regular index is built
    build_rag_index(limit=limit)
    
    init_late_interaction_db()
    
    print("[RAG Manager] Building Late Interaction index...", flush=True)
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Get list of already processed files in late interaction
    cursor.execute("SELECT DISTINCT file_path FROM reports_late_interaction")
    processed_files = {row[0] for row in cursor.fetchall()}
    
    # Get all reports text
    cursor.execute("SELECT file_path, content FROM reports_rag")
    all_reports = cursor.fetchall()
    conn.close()
    
    if not all_reports:
        print("[RAG Manager] No reports found in reports_rag to build Late Interaction index.", flush=True)
        return
        
    model = get_late_interaction_model()
    
    indexed_count = 0
    for file_path, content in all_reports:
        if file_path in processed_files:
            continue
            
        print(f"[RAG Manager] Chunking & embedding for ColBERT: {os.path.basename(file_path)}...", flush=True)
        chunks = chunk_text(content)
        if not chunks:
            continue
            
        try:
            # Get token embeddings for all chunks in a batch
            embeddings_list = model.encode(chunks, output_value="token_embeddings", batch_size=32)
            
            conn = sqlite3.connect(DB_PATH)
            db_cursor = conn.cursor()
            
            for idx, (chunk, emb) in enumerate(zip(chunks, embeddings_list)):
                emb_np = emb.cpu().numpy().astype(np.float32) if hasattr(emb, "cpu") else np.array(emb, dtype=np.float32)
                num_tokens = emb_np.shape[0]
                blob = emb_np.tobytes()
                
                db_cursor.execute("""
                    INSERT OR REPLACE INTO reports_late_interaction 
                    (file_path, chunk_index, chunk_text, token_embeddings_blob, num_tokens)
                    VALUES (?, ?, ?, ?, ?)
                """, (file_path, idx, chunk, blob, num_tokens))
                
            conn.commit()
            conn.close()
            indexed_count += 1
            print(f"[RAG Manager] Embedded {len(chunks)} chunks for {os.path.basename(file_path)}.", flush=True)
            
        except Exception as e:
            print(f"[RAG Manager] Error embedding chunks for {file_path}: {e}", flush=True)
            
        if limit is not None and indexed_count >= limit:
            break
            
    print(f"[RAG Manager] Late Interaction index updated. Processed {indexed_count} files.", flush=True)


def search_rag_late_interaction(query: str, limit: int = 5, file_path: Optional[str] = None) -> List[Dict[str, Any]]:
    """Query the ColBERT late-interaction index using token-level embeddings & MaxSim."""
    init_late_interaction_db()
    
    # 1. Embed query
    model = get_late_interaction_model()
    query_emb_list = model.encode([query], output_value="token_embeddings")
    if not query_emb_list:
        return []
    
    query_emb = query_emb_list[0]
    query_emb = query_emb.cpu().numpy().astype(np.float32) if hasattr(query_emb, "cpu") else np.array(query_emb, dtype=np.float32)
    
    # Slice out [CLS] and [SEP]
    if query_emb.shape[0] > 2:
        query_emb = query_emb[1:-1]
        
    q_norms = np.linalg.norm(query_emb, axis=1, keepdims=True)
    q_norms = np.where(q_norms == 0, 1e-9, q_norms)
    q_norm = query_emb / q_norms
    
    # 2. Fetch all chunks from DB
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    if file_path:
        # Standardize file path separators for comparison
        norm_filter = os.path.normpath(file_path)
        cursor.execute("""
            SELECT file_path, chunk_index, chunk_text, token_embeddings_blob, num_tokens 
            FROM reports_late_interaction
            WHERE file_path = ? OR file_path = ?
        """, (norm_filter, norm_filter.replace('\\', '/')))
    else:
        cursor.execute("""
            SELECT file_path, chunk_index, chunk_text, token_embeddings_blob, num_tokens 
            FROM reports_late_interaction
        """)
    rows = cursor.fetchall()
    conn.close()
    
    scored_results = []
    for file_path, chunk_index, chunk_text, blob, num_tokens in rows:
        if not blob or num_tokens <= 0:
            continue
            
        # Reconstruct document token embeddings
        d_emb = np.frombuffer(blob, dtype=np.float32).reshape((num_tokens, 384))
        
        # Slice out [CLS] and [SEP]
        if d_emb.shape[0] > 2:
            d_emb = d_emb[1:-1]
            
        d_norms = np.linalg.norm(d_emb, axis=1, keepdims=True)
        d_norms = np.where(d_norms == 0, 1e-9, d_norms)
        d_norm = d_emb / d_norms
        
        # Compute MaxSim
        sim_matrix = np.dot(q_norm, d_norm.T)
        max_sim = np.max(sim_matrix, axis=1)
        score = float(np.sum(max_sim))
        
        scored_results.append({
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "chunk_index": chunk_index,
            "chunk_text": chunk_text,
            "score": score
        })
        
    # Sort by score descending
    scored_results.sort(key=lambda x: x["score"], reverse=True)
    return scored_results[:limit]
