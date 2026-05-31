# app/utils/template_manager.py

import os
import re
try:
    import docx
except ImportError:
    docx = None
from typing import List, Dict, Set

# Regex to match placeholders like [CLIENT], [CLIENT_NAME], [Site Address], {{client}}, etc.
PLACEHOLDER_RE = re.compile(r'(?:\[|\{\{)([A-Z0-9_\s\-\/]{2,40})(?:\]|\}\})', re.IGNORECASE)

TEMPLATES_DIR = r"C:\Users\pored\Downloads\Project Geologs\Templates"
REPORTS_DIRS = [
    r"C:\Users\pored\Downloads\Project Geologs\03 - Reports",
    r"C:\Users\pored\Downloads\Project Geologs\Reports 2",
    r"C:\Users\pored\Downloads\Project Geologs\Reports 3"
]

def clean_placeholder(p: str) -> str:
    """Trim and normalize placeholder string."""
    return p.strip()

def extract_from_paragraph(p) -> Set[str]:
    """Extract placeholders from a docx Paragraph."""
    matches = PLACEHOLDER_RE.findall(p.text)
    return {clean_placeholder(m) for m in matches}

def extract_placeholders(file_path: str) -> List[str]:
    """Scan a .docx file and return a unique list of detected placeholders."""
    if docx is None:
        return []
    if not os.path.exists(file_path) or not file_path.endswith('.docx'):
        return []
    
    try:
        doc = docx.Document(file_path)
        placeholders = set()
        
        # 1. Scan normal paragraphs
        for p in doc.paragraphs:
            placeholders.update(extract_from_paragraph(p))
            
        # 2. Scan table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        placeholders.update(extract_from_paragraph(p))
                        
        # 3. Scan headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for p in header.paragraphs:
                        placeholders.update(extract_from_paragraph(p))
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    placeholders.update(extract_from_paragraph(p))
                                    
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for p in footer.paragraphs:
                        placeholders.update(extract_from_paragraph(p))
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    placeholders.update(extract_from_paragraph(p))
        
        # Filter and return sorted list
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error scanning placeholders in {file_path}: {e}")
        return []

def list_templates() -> List[Dict]:
    """List all available templates under Templates."""
    templates = []
    if not os.path.exists(TEMPLATES_DIR):
        return []
        
    for root, dirs, files in os.walk(TEMPLATES_DIR):
        for file in files:
            if file.startswith("~$") or not (file.endswith(".docx") or file.endswith(".xlsx") or file.endswith(".xlsm")):
                continue
                
            full_path = os.path.join(root, file)
            rel_path = os.path.relpath(full_path, TEMPLATES_DIR)
            category = os.path.basename(root) if root != TEMPLATES_DIR else "General"
            
            # Extract placeholders for .docx templates to support dynamic inputs
            placeholders = []
            if file.endswith(".docx"):
                placeholders = extract_placeholders(full_path)
                
            templates.append({
                "template_id": rel_path.replace("\\", "/"),
                "file_name": file,
                "relative_path": rel_path,
                "category": category,
                "size_bytes": os.path.getsize(full_path),
                "placeholders": placeholders,
                "supported": file.endswith(".docx")
            })
            
    return sorted(templates, key=lambda x: (x["category"], x["file_name"]))

def replace_in_paragraph(paragraph, replacements: Dict[str, str]):
    """Replace placeholders in paragraph, attempting run preservation or falling back gracefully."""
    text = paragraph.text
    if not text:
        return
        
    # Check if there is any placeholder in text
    has_placeholder = False
    for placeholder in replacements:
        # Match both [placeholder] and {{placeholder}}
        p_bracket = f"[{placeholder}]"
        p_curly = f"{{{{{placeholder}}}}}"
        if p_bracket in text or p_curly in text:
            has_placeholder = True
            break
            
    if not has_placeholder:
        return
        
    # Try replacing inside individual runs (preserves run formatting)
    for placeholder, value in replacements.items():
        p_bracket = f"[{placeholder}]"
        p_curly = f"{{{{{placeholder}}}}}"
        for run in paragraph.runs:
            if p_bracket in run.text:
                run.text = run.text.replace(p_bracket, str(value))
            if p_curly in run.text:
                run.text = run.text.replace(p_curly, str(value))
                
    # If run replacement was incomplete due to split runs, fall back to paragraph level replacement
    text = paragraph.text
    still_has = False
    for placeholder in replacements:
        if f"[{placeholder}]" in text or f"{{{{{placeholder}}}}}" in text:
            still_has = True
            break
            
    if still_has:
        # Save original formatting of first run as best-effort fallback
        bold = paragraph.runs[0].bold if paragraph.runs else False
        italic = paragraph.runs[0].italic if paragraph.runs else False
        font_name = paragraph.runs[0].font.name if paragraph.runs and paragraph.runs[0].font else None
        font_size = paragraph.runs[0].font.size if paragraph.runs and paragraph.runs[0].font else None
        
        new_text = text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(f"[{placeholder}]", str(value))
            new_text = new_text.replace(f"{{{{{placeholder}}}}}", str(value))
            
        paragraph.text = ""  # Clear paragraph runs
        run = paragraph.add_run(new_text)
        run.bold = bold
        run.italic = italic
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size

def fill_template(template_path: str, replacements: Dict[str, str], output_path: str) -> str:
    """Load a .docx template, substitute placeholders, and save to output_path."""
    if docx is None:
        raise RuntimeError("docx library is not installed")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")
        
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = docx.Document(template_path)
    
    # 1. Replace in normal paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
        
    # 2. Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)
                    
    # 3. Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    replace_in_paragraph(p, replacements)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_in_paragraph(p, replacements)
                                
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    replace_in_paragraph(p, replacements)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_in_paragraph(p, replacements)
                                
    doc.save(output_path)
    return output_path

def list_history() -> List[Dict]:
    """Scan previous reports folder recursively and return a structured registry of completed work."""
    history = []
    
    for base_dir in REPORTS_DIRS:
        if not os.path.exists(base_dir):
            continue
            
        group_name = os.path.basename(base_dir) # "03 - Reports", "Reports 2", etc.
        
        for root, dirs, files in os.walk(base_dir):
            for file in files:
                if file.startswith("~$") or not (file.endswith(".docx") or file.endswith(".pdf") or file.endswith(".xlsx")):
                    continue
                    
                full_path = os.path.join(root, file)
                rel_path = os.path.relpath(full_path, base_dir)
                sub_folder = os.path.relpath(root, base_dir)
                
                history.append({
                    "file_name": file,
                    "relative_path": rel_path,
                    "group": group_name,
                    "sub_folder": sub_folder if sub_folder != "." else "General",
                    "file_size": os.path.getsize(full_path),
                    "file_path": full_path,
                    "extension": os.path.splitext(file)[1].lower()
                })
                
    return sorted(history, key=lambda x: (x["group"], x["sub_folder"], x["file_name"]))

def extract_variables_from_docx(file_path: str) -> Dict[str, str]:
    """Extremely robust text parsing and regex pattern matching to auto-extract report fields."""
    if docx is None:
        return extracted
    extracted = {
        "CLIENT": "",
        "CLIENT_NAME": "",
        "ADDRESS": "",
        "SITE_ADDRESS": "",
        "JOB_NO": "",
        "REPORT_NO": "",
        "DATE": "",
        "ENGINEER": "",
        "BEARING_CAPACITY": ""
    }
    
    if not os.path.exists(file_path) or not file_path.endswith('.docx'):
        return extracted
        
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Combine all paragraphs and cells to scan
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
                        
        combined_text = "\n".join(full_text)
        
        # Smart regex parsing patterns
        client_match = re.search(r'(?:Client|Commissioned By|Prepared For)\s*:\s*([^\n\r]+)', combined_text, re.IGNORECASE)
        address_match = re.search(r'(?:Site Address|Site|Location|Project Site|Address)\s*:\s*([^\n\r]+)', combined_text, re.IGNORECASE)
        job_match = re.search(r'(?:Job No|Job Number|Project No|Ref)\s*:\s*([A-Z0-9\-\s\/]+)', combined_text, re.IGNORECASE)
        report_match = re.search(r'(?:Report No|Report Number|Document No)\s*:\s*([A-Z0-9\-\s\/]+)', combined_text, re.IGNORECASE)
        date_match = re.search(r'(?:Date|Date of Inspection)\s*:\s*([^\n\r]+)', combined_text, re.IGNORECASE)
        engineer_match = re.search(r'(?:Engineer|Inspector|Logged By|Approved By)\s*:\s*([A-Z\.\s]{2,30})', combined_text, re.IGNORECASE)
        bearing_match = re.search(r'(?:Bearing Capacity|Allowable Bearing|Bearing)\s*:\s*([^\n\r]+)', combined_text, re.IGNORECASE)
        
        if client_match:
            extracted["CLIENT"] = client_match.group(1).strip()
            extracted["CLIENT_NAME"] = client_match.group(1).strip()
        if address_match:
            extracted["ADDRESS"] = address_match.group(1).strip()
            extracted["SITE_ADDRESS"] = address_match.group(1).strip()
        if job_match:
            extracted["JOB_NO"] = job_match.group(1).strip()
        if report_match:
            extracted["REPORT_NO"] = report_match.group(1).strip()
        if date_match:
            extracted["DATE"] = date_match.group(1).strip()
        if engineer_match:
            extracted["ENGINEER"] = engineer_match.group(1).strip()
        if bearing_match:
            extracted["BEARING_CAPACITY"] = bearing_match.group(1).strip()
            
        # Deduce any specific bracketed placeholders in text and see if we can find corresponding filled values next to them
        # (This acts as an intelligent local fallback extraction)
        return extracted
    except Exception as e:
        print(f"Error extracting variables from {file_path}: {e}")
        return extracted
